#!/usr/bin/env python3
"""
Markdown → 党政机关公文格式 (.docx)
基于 GB/T 9704-2012《党政机关公文格式》实现，生成规范化 Word/WPS 文档。

Typora 一键导出（推荐）：
  偏好设置 → 导出 → 添加导出 → 自定义
  命令：python3 /完整路径/markdown-to-gongwen.py -d ~/Desktop

命令行：
  python3 markdown-to-gongwen.py 输入.md [输出.docx]
"""

from __future__ import annotations

import os
import sys
import subprocess

# ── 自动安装依赖（首次运行自动 pip install）─────────────
try:
    import regex
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("❌ 缺少依赖，请运行：.venv/bin/pip install python-docx regex", file=sys.stderr, flush=True)
    sys.exit(1)


# ============================================================
# 常量定义 —— 严格按 GB/T 9704-2012 设置
# ============================================================
class GongwenFormat:
    # 页面
    PAGE_WIDTH = Cm(21.0)          # A4
    PAGE_HEIGHT = Cm(29.7)
    TOP_MARGIN = Cm(3.7)           # 天头 37mm±1mm
    BOTTOM_MARGIN = Cm(3.5)
    LEFT_MARGIN = Cm(2.8)          # 订口 28mm±1mm
    RIGHT_MARGIN = Cm(2.6)

    # 正文字体字号
    BODY_FONT = "仿宋_GB2312"
    BODY_SIZE = Pt(16)             # 三号 = 16pt
    BODY_LINE_SPACING = Pt(28)     # 固定值 28 磅
    BODY_FIRST_LINE_INDENT = Pt(32)    # 三号字 16pt × 2 = 32pt（2 个中文字符）

    # 标题字体字号（文档主标题 —— 用“# ”标记）
    MAIN_TITLE_FONT = "方正小标宋简体"
    MAIN_TITLE_SIZE = Pt(22)       # 二号 = 22pt

    # 一级标题（## ）= 公文“一、”
    HEADING1_FONT = "黑体"
    HEADING1_SIZE = Pt(16)         # 三号

    # 二级标题（### ）= 公文“（一）”
    HEADING2_FONT = "楷体_GB2312"
    HEADING2_SIZE = Pt(16)         # 三号

    # 三级标题（#### ）= 公文 “1.”
    HEADING3_FONT = "仿宋_GB2312"
    HEADING3_SIZE = Pt(16)         # 三号，加粗

    # 四级标题（##### 及更深）= 公文“（1）”
    HEADING4_FONT = "仿宋_GB2312"
    HEADING4_SIZE = Pt(16)         # 三号

    # 数字/英文字体
    WESTERN_FONT = "Times New Roman"


# ============================================================
# 全角半角处理工具
# ============================================================
def fullwidth_to_halfwidth(text: str) -> str:
    """将全角字母/数字转为半角（不影响全角中文标点）。"""
    # 全角→半角偏移量
    OFFSET = 0xFEE0
    result = []
    for ch in text:
        code = ord(ch)
        if code == 0x3000:                # 全角空格
            result.append(' ')
        elif 0xFF10 <= code <= 0xFF19:    # 全角数字 ０-９
            result.append(chr(code - OFFSET))
        elif 0xFF21 <= code <= 0xFF3A:    # 全角大写字母 Ａ-Ｚ
            result.append(chr(code - OFFSET))
        elif 0xFF41 <= code <= 0xFF5A:    # 全角小写字母 ａ-ｚ
            result.append(chr(code - OFFSET))
        else:
            result.append(ch)
    return ''.join(result)
def halfwidth_punct_to_fullwidth(text: str) -> str:
    """自动识别上下文转换标点：中文环境使用全角，英文/数字环境使用半角。"""

    def is_cjk(ch: str) -> bool:
        """判断是否为中文字符（CJK 统一表意文字及CJK符号）。"""
        cp = ord(ch)
        return (0x4E00 <= cp <= 0x9FFF) or (0x3000 <= cp <= 0x303F)

    mapping = {
        ',': '，', ';': '；', ':': '：',
        '?': '？', '!': '！',
        '(': '（', ')': '）', '[': '［', ']': '］',
    }

    result = list(text)
    for i, ch in enumerate(result):
        if ch not in mapping and ch != '.':
            continue

        prev = result[i - 1] if i > 0 else ''
        next_ch = result[i + 1] if i < len(result) - 1 else ''

        if ch == '.':
            # 句点：数字编号/版本号保持半角，中文环境转全角
            if prev and (prev.isdigit() or prev.isalpha()):
                # "1." 编号、"3.0" 版本号 → 半角
                if (not next_ch) or next_ch.isspace() or next_ch.isdigit() or next_ch.isalpha():
                    continue
            # 前后有CJK → 全角句号
            if is_cjk(prev) or is_cjk(next_ch):
                result[i] = '。'
            # 孤立的句点 → 默认全角（中文文档）
            elif not (prev or next_ch):
                result[i] = '。'
            continue

        # 其他标点：CJK 上下文 → 全角
        if is_cjk(prev) or is_cjk(next_ch):
            result[i] = mapping[ch]
        # 否则保持半角（纯英文/数字环境）

    return ''.join(result)


def ensure_liu_jiao_bracket(text: str) -> str:
    """
    确保发文字号类型的括号为六角〔〕，替代 []。
    匹配模式：类似 [2026] 7号
    """
    pattern = regex.compile(r'\[(\d{4})\](?=\s*\d+\s*号)')
    return pattern.sub(r'〔\1〕', text)


def strip_markdown_markers(text: str) -> str:
    """去除正文中的 Markdown 格式标记（列表符号、加粗）。"""
    # 去除段首无序列表标记：- 或 *
    text = regex.sub(r'^[-*]\s+', '', text)
    # 去除 **加粗** 标记，保留内部文字
    text = regex.sub(r'\*\*(.+?)\*\*', r'\1', text)
    return text


# ============================================================
# 段落格式辅助
# ============================================================
def set_run_font(run, font_name: str, size: Pt, bold: bool = False, western: str = ""):
    """设置 run 的中文字体、西文字体和字号。"""
    run.bold = bold
    run.font.size = size
    run.font.name = font_name
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r.insert(0, rPr)
    # 中文字体
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    if western:
        rFonts.set(qn('w:ascii'), western)
        rFonts.set(qn('w:hAnsi'), western)


def set_paragraph_spacing(paragraph, line_spacing: Pt, first_line_indent=None,
                          alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                          space_after=Pt(0)):
    """设置段落行距、首行缩进、段后间距和对齐方式。"""
    pf = paragraph.paragraph_format
    pf.line_spacing = line_spacing
    pf.alignment = alignment
    pf.space_after = space_after
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent


def add_formatted_paragraph(doc, text: str, font_name: str, size: Pt,
                            bold: bool = False, first_line_indent=None,
                            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                            western: str = "", space_after=Pt(0)) -> "Paragraph":
    """向文档添加一个格式化好的段落，并返回该段落对象。"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, font_name, size, bold, western)
    set_paragraph_spacing(p, GongwenFormat.BODY_LINE_SPACING,
                          first_line_indent, alignment, space_after)
    return p


def apply_western_font_to_digits_and_letters(paragraph, western: str = "Times New Roman"):
    """遍历段落中所有 run，将数字和字母部分单独设置为西文字体。"""
    for run in paragraph.runs:
        original_text = run.text
        if not original_text:
            continue
        # 查找数字/英文字母片段并分割
        parts = regex.split(r'([A-Za-z0-9]+)', original_text)
        if len(parts) <= 1:
            continue
        # 清除原有 run 文本并重建多个 run
        run.text = ""
        for part in parts:
            if not part:
                continue
            new_run = paragraph.add_run(part)
            if regex.match(r'[A-Za-z0-9]+', part):
                set_run_font(new_run, GongwenFormat.BODY_FONT,
                             GongwenFormat.BODY_SIZE, western=western)
            else:
                set_run_font(new_run, GongwenFormat.BODY_FONT,
                             GongwenFormat.BODY_SIZE)
        # 移除旧的空 run
        paragraph._element.remove(run._element)


# ============================================================
# Markdown 解析与层级映射
# ============================================================
class MarkdownBlock:
    """Markdown 文本块（标题或普通段落）。"""
    def __init__(self, line: str):
        self.raw = line.strip()
        self._parse()

    def _parse(self):
        self.type = "paragraph"
        self.level = 0
        self.text = ""
        if regex.match(r'^#{1,6}\s', self.raw):
            self.type = "heading"
            m = regex.match(r'^(#+)\s*(.*)', self.raw)
            self.level = len(m.group(1))
            self.text = m.group(2).strip()
        else:
            self.text = self.raw


def apply_punctuation_norm(text: str) -> str:
    """把正文文本的标点符号规范化。"""
    # 先去除 Markdown 标记（列表符号、加粗等）
    text = strip_markdown_markers(text)
    text = halfwidth_punct_to_fullwidth(text)
    text = fullwidth_to_halfwidth(text)
    text = ensure_liu_jiao_bracket(text)
    return text


# ============================================================
# 主转换函数
# ============================================================
def markdown_to_gongwen(
    markdown_text: str,
    output_path: str = "output_gongwen.docx",
) -> str:
    """
    将 Markdown 文本转换为符合 GB/T 9704-2012 公排版式的 .docx 文件。
    返回生成文件的绝对路径。
    """

    # ----- 1. 解析 Markdown -----
    lines = markdown_text.splitlines()
    blocks = [MarkdownBlock(line) for line in lines if line.strip() != ""]

    main_title_text = "公文标题"
    heading1_level_counter = 0
    date_string = None
    zhusong_list: list[str] = []

    # 预扫描：提取主标题、主送机关、成文日期
    for block in blocks:
        if block.type == "heading" and block.level == 1 and main_title_text == "公文标题":
            main_title_text = block.text.strip()
        if block.type == "paragraph":
            # 主送机关识别：包含“主送机关：”字样
            if block.text.startswith("主送机关：") or block.text.startswith("主送机关:"):
                raw_recipients = block.text.replace("主送机关：", "").replace("主送机关:", "").strip()
                zhusong_list = [r.strip() for r in raw_recipients.split("、") if r.strip()]
                blocks.remove(block)
                continue
            # 成文日期识别：包含“成文日期：”字样
            if block.text.startswith("成文日期：") or block.text.startswith("成文日期:"):
                raw_date = block.text.replace("成文日期：", "").replace("成文日期:", "").strip()
                date_string = raw_date
                blocks.remove(block)
                continue

    # 去除已特殊处理的 blocks
    blocks = [b for b in blocks if not (
        (b.text.startswith("主送机关") and b.level == 0) or
        (b.text.startswith("成文日期") and b.level == 0)
    )]

    # 重新查找主标题（去除特殊要素后）
    for block in blocks:
        if block.type == "heading" and block.level == 1:
            main_title_text = block.text.strip()
            break

    # ----- 2. 构建 Word 文档 -----
    doc = Document()

    # --- 页面设置 ---
    section = doc.sections[0]
    section.page_width = GongwenFormat.PAGE_WIDTH
    section.page_height = GongwenFormat.PAGE_HEIGHT
    section.top_margin = GongwenFormat.TOP_MARGIN
    section.bottom_margin = GongwenFormat.BOTTOM_MARGIN
    section.left_margin = GongwenFormat.LEFT_MARGIN
    section.right_margin = GongwenFormat.RIGHT_MARGIN

    # --- 主标题 ---
    title_p = add_formatted_paragraph(
        doc, main_title_text,
        GongwenFormat.MAIN_TITLE_FONT,
        GongwenFormat.MAIN_TITLE_SIZE,
        bold=False,
        first_line_indent=None,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    # 主标题后空一行
    blank = doc.add_paragraph()
    set_paragraph_spacing(blank, GongwenFormat.BODY_LINE_SPACING, space_after=Pt(0))

    # --- 主送机关 ---
    if zhusong_list:
        recipient_text = "、".join(zhusong_list) + "："
        add_formatted_paragraph(
            doc, recipient_text,
            GongwenFormat.BODY_FONT,
            GongwenFormat.BODY_SIZE,
            first_line_indent=None,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
        )

    # --- 正文/标题层级 ---
    heading1_counter = 0
    for block in blocks:
        if block.type == "paragraph":
            raw_text = block.text
            # 跳过已处理的主送机关和成文日期
            if raw_text.startswith("主送机关") or raw_text.startswith("成文日期"):
                continue
            norm_text = apply_punctuation_norm(raw_text)
            # 段落
            p = add_formatted_paragraph(
                doc, norm_text,
                GongwenFormat.BODY_FONT,
                GongwenFormat.BODY_SIZE,
                first_line_indent=GongwenFormat.BODY_FIRST_LINE_INDENT,
            )
            apply_western_font_to_digits_and_letters(p, GongwenFormat.WESTERN_FONT)

        elif block.type == "heading":
            if block.level == 1:
                # 主标题已经输出，跳过
                continue
            elif block.level == 2:
                heading1_counter += 1
                # 二级标题（公文一级标题"一、"）：缩进2字符
                cns = ['一', '二', '三', '四', '五', '六', '七', '八', '九', '十']
                prefix = f"{cns[heading1_counter - 1]}、" if heading1_counter <= len(cns) else ""
                level_text = f"{prefix}{block.text}"
                p = add_formatted_paragraph(
                    doc, level_text,
                    GongwenFormat.HEADING1_FONT,
                    GongwenFormat.HEADING1_SIZE,
                    bold=True,
                    first_line_indent=GongwenFormat.BODY_FIRST_LINE_INDENT,
                    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                )
            elif block.level == 3:
                norm_text = apply_punctuation_norm(block.text)
                p = add_formatted_paragraph(
                    doc, f"（一）{norm_text}" if not norm_text.startswith("（") else norm_text,
                    GongwenFormat.HEADING2_FONT,
                    GongwenFormat.HEADING2_SIZE,
                    first_line_indent=GongwenFormat.BODY_FIRST_LINE_INDENT,  # 🔥 新增缩进
                    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                )
            elif block.level == 4:
                norm_text = apply_punctuation_norm(block.text)
                p = add_formatted_paragraph(
                    doc, f"1. {norm_text}" if not norm_text.startswith("1.") else norm_text,
                    GongwenFormat.HEADING3_FONT,
                    GongwenFormat.HEADING3_SIZE,
                    bold=True,
                    first_line_indent=GongwenFormat.BODY_FIRST_LINE_INDENT,  # 🔥 新增缩进
                )
            elif block.level >= 5:
                norm_text = apply_punctuation_norm(block.text)
                p = add_formatted_paragraph(
                    doc, f"（1）{norm_text}" if not norm_text.startswith("（") else norm_text,
                    GongwenFormat.HEADING4_FONT,
                    GongwenFormat.HEADING4_SIZE,
                    first_line_indent=GongwenFormat.BODY_FIRST_LINE_INDENT,  # 🔥 新增缩进
                )
            # 对标题也校正西文字体
            if block.level >= 2:
                apply_western_font_to_digits_and_letters(p, GongwenFormat.WESTERN_FONT)
    # --- 落款（成文日期）---
    if date_string:
        # 空两行
        doc.add_paragraph()
        doc.add_paragraph()
        date_p = doc.add_paragraph()
        date_run = date_p.add_run(date_string)
        set_run_font(date_run, GongwenFormat.BODY_FONT, GongwenFormat.BODY_SIZE)
        set_paragraph_spacing(date_p, GongwenFormat.BODY_LINE_SPACING,
                              alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        # 右空四字：通过右侧缩进实现
        date_p.paragraph_format.right_indent = Cm(1.7)

    # ----- 3. 保存文件 -----
    doc.save(output_path)
    return os.path.abspath(output_path)


# ============================================================
# 便捷使用函数
# ============================================================
def convert_markdown_to_gongwen(
    markdown_content: str,
    output_filename: str = "公文输出.docx",
) -> str:
    """
    用户友好接口：输入 Markdown 内容字符串，返回生成 docx 文件完整路径。
    """
    return markdown_to_gongwen(markdown_content, output_filename)


def _detect_typora_document() -> str | None:
    """macOS: 探测 Typora 当前文档路径（无需 Accessibility 权限）。"""

    def _via_axdocument() -> str | None:
        """方案A：通过 AXDocument 属性（需 Accessibility 权限）。"""
        try:
            cmd = [
                "osascript", "-e",
                'tell application "System Events" to tell process "Typora" '
                'to get value of attribute "AXDocument" of front window'
            ]
            r = subprocess.run(cmd, capture_output=True, text=True, timeout=5)
            if r.returncode == 0 and r.stdout.strip():
                url = r.stdout.strip()
                if url.startswith("file://"):
                    from urllib.parse import unquote, urlparse
                    path = unquote(urlparse(url).path)
                    if os.path.exists(path):
                        return path
        except Exception:
            pass
        return None

    def _via_window_title() -> str | None:
        """方案B：通过窗口标题搜索文件（无需额外权限）。"""
        try:
            cmd = [
                "osascript", "-e",
                'tell application "Typora" to get name of window 1'
            ]
            r = subprocess.run(cmd, capture_output=True, text=True, timeout=5)
            if r.returncode != 0 or not r.stdout.strip():
                return None
            win_title = r.stdout.strip()
            if not win_title or win_title.startswith("未命名"):
                return None

            base = win_title[:-3] if win_title.endswith(".md") else win_title
            fname = base + ".md"

            search_dirs = [
                os.path.expanduser("~/Documents"),
                os.path.expanduser("~/Desktop"),
                os.getcwd(),
            ]
            for d in search_dirs:
                full = os.path.join(d, fname)
                if os.path.exists(full) and os.path.isfile(full):
                    return os.path.abspath(full)

            sp = subprocess.run(
                ["mdfind", "-onlyin", os.path.expanduser("~"),
                 f"kMDItemFSName == '{fname}'"],
                capture_output=True, text=True, timeout=10
            )
            for line in sp.stdout.splitlines():
                p = line.strip()
                if os.path.exists(p) and p.endswith(".md"):
                    return p
        except Exception:
            pass
        return None

    return _via_axdocument() or _via_window_title()


# ============================================================
# 命令行入口
# ============================================================
if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(
        description="Markdown 转党政机关公文格式 (.docx)"
    )
    parser.add_argument(
        "input", nargs="?",
        help="输入 Markdown 文件路径（不传则自动探测 Typora 当前文档）"
    )
    parser.add_argument(
        "output_legacy", nargs="?",
        help=argparse.SUPPRESS  # 向后兼容：python script.py input.md output.docx
    )
    parser.add_argument(
        "-o", "--output", default=None,
        help="输出 .docx 文件路径"
    )
    parser.add_argument(
        "-d", "--dir", default=None,
        help="输出目录（默认与输入文件同目录，如 -d ~/Desktop）"
    )
    args = parser.parse_args()

    md = None
    input_filepath = None

    # ── 模式A：传了文件路径（终端命令行模式） ──
    if args.input:
        input_filepath = args.input.strip("\"'")
        if not os.path.exists(input_filepath):
            print(f"❌ 找不到文件: {input_filepath}", file=sys.stderr)
            sys.exit(1)
        with open(input_filepath, "r", encoding="utf-8") as f:
            md = f.read()

    # ── 模式B：探测 Typora 当前文档 ──
    if md is None:
        input_filepath = _detect_typora_document()
        if input_filepath:
            with open(input_filepath, "r", encoding="utf-8") as f:
                md = f.read()

    # ── 模式C：stdin 兜底 ──
    if md is None:
        md = sys.stdin.read()

    if not md or not md.strip():
        print("❌ 错误：未读取到 Markdown 内容", file=sys.stderr)
        print("   用法：python3 markdown-to-gongwen.py 输入.md", file=sys.stderr)
        print("   或 Typora 导出：python3 markdown-to-gongwen.py -d ~/Desktop", file=sys.stderr)
        sys.exit(1)

    # ── 确定输出路径 ──
    if args.output:
        output_path = os.path.expanduser(args.output)
    elif args.output_legacy:
        output_path = os.path.expanduser(args.output_legacy)
    elif args.dir or input_filepath:
        if args.dir:
            out_dir = os.path.expanduser(args.dir)
        else:
            out_dir = os.path.dirname(input_filepath) if input_filepath else "."
        os.makedirs(out_dir, exist_ok=True)
        base = os.path.splitext(
            os.path.basename(input_filepath) if input_filepath else "公文输出"
        )[0]
        output_path = os.path.join(out_dir, base + ".docx")
    else:
        output_path = "公文输出.docx"

    result_path = convert_markdown_to_gongwen(md, output_path)
    print(f"✅ 公文已生成: {result_path}")